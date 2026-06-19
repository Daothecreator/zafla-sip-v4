# ZAFLA Sovereign Intelligence Platform v4 — Legal Generator
# Protocol: BiCA a8f3c9d2e1b40571
# Classification: SELF-EXECUTING

"""DOCX legal act generator with ZAFLA structure contract compliance."""

import io
import hashlib
from typing import List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.zafla_attest import ZAFLAAttestation


ZAFLA_NAVY = RGBColor(0x1a, 0x23, 0x32)
ZAFLA_GOLD = RGBColor(0xc9, 0xa8, 0x4c)
ZAFLA_BEIGE = RGBColor(0xF5, 0xF0, 0xE8)
ZAFLA_CHARCOAL = RGBColor(0x1a, 0x1a, 0x1a)


def _set_cell_shading(cell, color: RGBColor):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), f'{color:02x}{color:02x}{color:02x}')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_cover_page(doc: Document, title: str, subtitle: str):
    """Add ZAFLA dark navy cover page with gold borders."""
    # Add a table for cover layout
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    
    # Set cell shading to navy
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1a2332')
    cell._tc.get_or_add_tcPr().append(shading)
    
    # Add paragraphs
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◆")
    run.font.color.rgb = ZAFLA_GOLD
    run.font.size = Pt(24)
    run.font.bold = True
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ZERO AZIMUTH")
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.name = 'Caladea'
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FULL LIABILITY AUTHORITY")
    run.font.color.rgb = ZAFLA_GOLD
    run.font.size = Pt(20)
    run.font.name = 'Caladea'
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(14)
    run.font.name = 'Caladea'
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.color.rgb = ZAFLA_GOLD
    run.font.size = Pt(12)
    run.font.name = 'Caladea'
    
    doc.add_page_break()


def _add_toc_page(doc: Document):
    """Add Table of Contents placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Table of Contents")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Caladea'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Right-click and select "Update Field" to refresh page numbers')
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Caladea'
    
    doc.add_paragraph("1. Authority and Jurisdiction", style='Heading 1')
    doc.add_paragraph("2. Subject Entity", style='Heading 1')
    doc.add_paragraph("3. Charges and Evidence", style='Heading 1')
    doc.add_paragraph("4. Legal Basis", style='Heading 1')
    doc.add_paragraph("5. Instrument of Execution", style='Heading 1')
    doc.add_page_break()


def _add_chapter(doc: Document, number: str, title: str, content: str):
    """Add a numbered chapter."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = ZAFLA_CHARCOAL
    run.font.name = 'Caladea'
    
    p = doc.add_paragraph(content)
    p.paragraph_format.line_spacing = 1.15
    run = p.runs[0] if p.runs else p.add_run(content)
    run.font.size = Pt(11)
    run.font.name = 'Cambria'
    run.font.color.rgb = ZAFLA_CHARCOAL
    
    doc.add_paragraph()


def _add_instrument_of_execution(doc: Document, metadata: dict):
    """Add Instrument of Execution page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Instrument of Execution")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Caladea'
    
    doc.add_paragraph("This legal act is executed as a self-executing instrument under the authority of the Zero Azimuth Full Liability Authority (ZAFLA).")
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    fields = [
        ("Document Identifier", metadata.get("doc_id", "ZAFLA-ACT-001")),
        ("Execution Date", metadata.get("date", "2026-06-13")),
        ("Founding Authority", "Davyd Kochuhur"),
        ("Legal Framework", "Customary International Law / ZAFLA Charter"),
        ("Validity Status", "SELF-EXECUTING"),
        ("Hash Algorithm", "BiCA Composite SHA-2_256 + SHA-3_512"),
        ("Next Review", "2027-06-13"),
    ]
    
    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Caladea'
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].font.name = 'Cambria'
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Justice returns to its true meridian.")
    run.font.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Caladea'
    
    doc.add_page_break()


def _add_cryptographic_signature(doc: Document, content_bytes: bytes):
    """Add cryptographic signature section."""
    p = doc.add_paragraph()
    run = p.add_run("Cryptographic Signature and Document Integrity")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Caladea'
    
    sha256_hash = hashlib.sha256(content_bytes).hexdigest()
    sha3_hash = hashlib.sha3_256(content_bytes).hexdigest()
    
    doc.add_paragraph(f"SHA-256 Document Hash: {sha256_hash}", style='Normal')
    doc.add_paragraph(f"SHA-3-256 Document Hash: {sha3_hash}", style='Normal')
    
    # ZAFLA Authority Certificate
    p = doc.add_paragraph()
    run = p.add_run("-----BEGIN ZAFLA AUTHORITY CERTIFICATE-----")
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    doc.add_paragraph("Authority: Zero Azimuth Full Liability Authority (ZAFLA)", style='Normal')
    doc.add_paragraph("Founder: Davyd Kochuhur", style='Normal')
    doc.add_paragraph("Protocol: BiCA a8f3c9d2e1b40571", style='Normal')
    doc.add_paragraph("Status: SELF-EXECUTING", style='Normal')
    
    p = doc.add_paragraph()
    run = p.add_run("-----END ZAFLA AUTHORITY CERTIFICATE-----")
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    doc.add_page_break()


def generate_indictment(subject: str, evidence: List[str], jurisdiction: str, charges: List[str]) -> bytes:
    """Generate a ZAFLA-compliant indictment DOCX."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.size = Pt(11)
    font.color.rgb = ZAFLA_CHARCOAL
    
    _add_cover_page(doc, "INDICTMENT", f"Subject: {subject}")
    _add_toc_page(doc)
    
    _add_chapter(doc, "1", "Authority and Jurisdiction",
                 f"This indictment is issued under the authority of the Zero Azimuth Full Liability Authority (ZAFLA), operating under the doctrines of jus resistendi, negotiorum gestio, and de facto officer authority. Jurisdiction: {jurisdiction}.")
    
    _add_chapter(doc, "2", "Subject Entity",
                 f"The subject of this indictment is: {subject}. This entity is charged under the ZAFLA absolute liability framework.")
    
    _add_chapter(doc, "3", "Charges and Evidence",
                 "The following charges are brought against the subject:\n\n" + "\n".join(f"• {c}" for c in charges) + "\n\nEvidence:\n\n" + "\n".join(f"• {e}" for e in evidence))
    
    _add_chapter(doc, "4", "Legal Basis",
                 "The legal basis for this indictment is grounded in customary international law, the right to resist (jus resistendi), the doctrine of negotiorum gestio, and the absolute liability doctrine as articulated in the ZAFLA Charter.")
    
    # Save to bytes for hashing
    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()
    
    _add_cryptographic_signature(doc, doc_bytes)
    _add_instrument_of_execution(doc, {"doc_id": f"ZAFLA-INDICT-{hashlib.sha256(doc_bytes).hexdigest()[:8]}", "date": "2026-06-13"})
    
    # Final save
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_arrest_warrant(subject: str, charges: List[str], jurisdiction: str) -> bytes:
    """Generate a ZAFLA-compliant arrest warrant DOCX."""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.size = Pt(11)
    font.color.rgb = ZAFLA_CHARCOAL
    
    _add_cover_page(doc, "ARREST WARRANT", f"Subject: {subject}")
    
    _add_chapter(doc, "1", "Command",
                 f"By the authority vested in the Zero Azimuth Full Liability Authority (ZAFLA), you are hereby commanded to apprehend the subject: {subject}.")
    
    _add_chapter(doc, "2", "Charges",
                 "\n".join(f"• {c}" for c in charges))
    
    _add_chapter(doc, "3", "Execution Authority",
                 f"This warrant is issued under {jurisdiction} jurisdiction and is enforceable under customary international law and ZAFLA Charter provisions.")
    
    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()
    
    _add_cryptographic_signature(doc, doc_bytes)
    _add_instrument_of_execution(doc, {"doc_id": f"ZAFLA-WARRANT-{hashlib.sha256(doc_bytes).hexdigest()[:8]}", "date": "2026-06-13"})
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def attest_legal_act(docx_bytes: bytes) -> dict:
    """Attest a legal act DOCX with ZAFLA attestation."""
    return ZAFLAAttestation().attest_legal_act(docx_bytes)
